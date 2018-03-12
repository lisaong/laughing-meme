#!/usr/bin/python3

import argparse
import glob
import itertools
import json
import os

# pip install python-docx
import docx

def jsonify_tscs(root_path, json_outfile, json_format):
    """Extracts TSCs from multiple docx files into a json file"""

    class Tsc:
        "Encapsulates a TSC"
        def __init__(self, tables):
            self.tables = tables
            self.row_offsets = []
            self.name = None
            self.category = None
            self.description = None
            self.num_levels = 0
            self.proficiencies = []
            self._extract()

        def _row_data(self, row):
            "returns row data, or empty columns if row is not defined"
            if row:
                HEADER_COLUMN = 0
                return [cell.text for cell in row.cells[HEADER_COLUMN+1:]]
            return [""] * self.num_levels

        def _row(self, index):
            # Initialize the offsets
            if not self.row_offsets:
                offset = 0
                for table in self.tables:
                    offset = offset + len(table.rows)
                    self.row_offsets.append(offset)

            prev_offset = 0
            for i, offset in enumerate(self.row_offsets):
                if index < offset:
                    return self.tables[i].rows[index - prev_offset]
                prev_offset = offset

            print("WARNING: index {} out of range, treating as empty row".format(index))
            return None

        def _extract(self):
            "Extracts a TSC from a docx table"

            # Row indices in the docx table
            INDICES = ["category", "name", "description",
                "prof.level", "prof.code", "prof.desc",
                "prof.knowledge",  "prof.abilities", "prof.apps"]

            # drop the repeated category cells for these rows
            self.name = self._row_data(self._row(INDICES.index("name")))[0]
            self.description = self._row_data(self._row(INDICES.index("description")))[0]
            self.category = self._row_data(self._row(INDICES.index("category")))[0]

            # proficiency rows
            prof_level = self._row_data(self._row(INDICES.index("prof.level")))
            self.num_levels = len(prof_level)
            prof_code = self._row_data(self._row(INDICES.index("prof.code")))
            prof_desc = self._row_data(self._row(INDICES.index("prof.desc")))
            prof_knowledge = self._row_data(self._row(INDICES.index("prof.knowledge")))
            prof_abilities = self._row_data(self._row(INDICES.index("prof.abilities")))
            prof_apps = self._row_data(self._row(INDICES.index("prof.apps")))

            # match naming to the document
            self.proficiencies = [{ "level" : level, "TSC code" : code, \
                "TSC Proficiency Description" : desc, "Knowledge" : knowledge, \
                "Abilities" : abilities, "Range of Applications" : apps} \
                    for level, code, desc, knowledge, abilities, apps in \
                        zip(prof_level, prof_code, prof_desc, prof_knowledge,
                            prof_abilities, prof_apps)
            ]

    class TscEncoder(json.JSONEncoder):
        "Extends JSON encoding for a TSC"
        # Suppress false positive pylint "hidden" warning
        def default(self, obj):  # pylint: disable=E0202
            if isinstance(obj, Tsc):
                return {
                    # match naming to the document
                    "Name" : obj.name,
                    "Category" : obj.category,
                    "Description" : obj.description,
                    "Proficiencies" : obj.proficiencies
                }

            # Let the base class default method raise the TypeError
            return json.JSONEncoder.default(self, obj)

    def flatten_tscs(tscs):
        "Flattens a list of TSCs into proficiencies"
        return list(itertools.chain.from_iterable(
            [[{
                "TSC Name" : tsc.name,
                "TSC Category" : tsc.category,
                "TSC Description" : tsc.description,
                **p} for p in tsc.proficiencies]
            for tsc in tscs]
        ))

    def get_top_level_tables(path):
        "Returns the top level tables in the document, if any"
        try:
            doc = docx.Document(path)
            if doc.tables:
                print(path)
                return doc.tables
        except:
            print("WARNING: Skipping invalid document {}".format(path))
        return None

    # main logic
    tscs = []
    for path in glob.glob("{}/**/*.docx".format(root_path),
        recursive=True):
        tables = get_top_level_tables(path)
        if tables:
            tscs.append(Tsc(tables))

    with open(json_outfile, mode='w') as o:
        if json_format == "tabular":
            tscs = flatten_tscs(tscs)
        o.write(json.dumps(tscs, cls=TscEncoder, indent=4,
            ensure_ascii=True))

if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description='Converts the top level table in a .docx to JSON.')
    parser.add_argument('path',
        help='folder path to look for .docx files')
    parser.add_argument('--outfile',
        help="output file path (default: skillsmap_table.json)",
        default="skillsmap_table.json")
    parser.add_argument('--format',
        help="determines the json output format (default: tabular).\n"
        "nested: grouped by TSCs, tabular: flat list of proficiencies",
        choices=["nested", "tabular"],
        default="tabular")
    args = parser.parse_args()
    jsonify_tscs(args.path, args.outfile, args.format)
